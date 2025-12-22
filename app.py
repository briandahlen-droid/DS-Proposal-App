"""
Development Services - Master Services Agreement IPO Generator
Streamlit web application for generating Individual Project Order documents
Complete single-file application
"""

import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

# ============================================================================
# TASK DESCRIPTIONS DATABASE
# ============================================================================

DEFAULT_FEES = {
    '110': {'name': 'Civil Engineering Design', 'amount': 40000, 'type': 'Hourly, Not-to-Exceed'},
    '120': {'name': 'Civil Schematic Design', 'amount': 35000, 'type': 'Hourly, Not-to-Exceed'},
    '130': {'name': 'Civil Design Development', 'amount': 45000, 'type': 'Hourly, Not-to-Exceed'},
    '140': {'name': 'Civil Construction Documents', 'amount': 50000, 'type': 'Hourly, Not-to-Exceed'},
    '150': {'name': 'Civil Permitting', 'amount': 40000, 'type': 'Hourly, Not-to-Exceed'},
    '210': {'name': 'Meetings and Coordination', 'amount': 20000, 'type': 'Hourly, Not-to-Exceed'}
}

TASK_DESCRIPTIONS = {
    '110': [
        "Kimley-Horn will prepare an onsite drainage report with supporting calculations showing the proposed development plan is consistent with the Southwest Florida Water Management District Basis of Review. This design will account for the stormwater design to support the development of the project site. The drainage report will include limited stormwater modeling to demonstrate that the Lot A site development will maintain the existing discharge rate and provide the required stormwater attenuation.",
        "The onsite drainage report will include calculations for 25-year 24-hour and 100-year 24-hour design storm conditions in accordance with Southwest Florida Water Management District Guidelines. A base stormwater design will be provided for the project site showing reasonable locations for stormwater conveyance features and stormwater management pond sizing."
    ],
    '120': [
        "Kimley-Horn will prepare Civil Schematic Design deliverables in accordance with the Client's Design Project Deliverables Checklist. For the Civil Schematic Design task, the deliverables that Kimley-Horn will provide consist of Civil Site Plan, Establish Finish Floor Elevations, Utility Will Serve Letters and Points of Service, Utility Routing and Easement Requirements."
    ],
    '130': [
        "Upon Client approval of the Schematic Design task, Kimley-Horn will prepare Design Development Plans of the civil design in accordance with the Client's Design Project Deliverables Checklist for Civil Design Development Deliverables. These documents will be approximately 50% complete and will include detail for City code review and preliminary pricing but will not include enough detail for construction bidding."
    ],
    '140': [
        "Based on the approved Development Plan, Kimley-Horn will provide engineering and design services for the preparation of site construction plans for on-site improvements.",
        "Cover Sheet",
        "The cover sheet includes plan contents, vicinity map, legal description and team identification.",
        "Existing Conditions Plan/Demolition Plan",
        "This sheet will include and identify the required demolition of the existing items on the project site.",
        "Site Layout Plan",
        "This sheet will include building setback lines, property lines, outline of building footprint, parking areas, handicap access ramps, sidewalks, crosswalks, driveways, and traffic lanes.",
        "Grading and Drainage Plan",
        "This sheet will include existing and proposed spot elevations and contours, building finish floor elevations, parking area drainage patterns, and stormwater inlet and pipe locations and sizes.",
        "Utility Plan",
        "This sheet will show the location and size of all water, sanitary sewer and reclaimed water facilities required to serve the development.",
        "Erosion and Sediment Control Plan",
        "This sheet will include erosion and sediment control measures designed to be implemented during construction.",
        "Details",
        "Standard and modified typical construction details will be provided."
    ],
    '150': [
        "Prepare and submit on the Client's behalf the following permitting packages for review/approval of construction documents, and attend meetings required to obtain the following Agency approvals:",
        "USF Site Development Permit",
        "Southwest Florida Water Management District Environmental Resource Permit – Minor Modification",
        "City of Tampa Water Department Commitment / Construction Plan Approval",
        "Hillsborough County Environmental Protection Commission",
        "Kimley-Horn will coordinate with the City of Tampa Development Review and coordination with the Florida Department of Transportation and the Hillsborough County departments as needed to obtain the necessary regulatory and utility approval of the site plans and associated drainage facilities. We will assist the Client with meetings necessary to gain site plan approval.",
        "This scope does not anticipate a Geotechnical or Environmental Assessment Report, Survey, Topographic Survey, or Arborist Report be required for this permit application.",
        "It is assumed Client will provide the needed information regarding the development program and requirements. Kimley-Horn will work with the Owner and their team to integrate the necessary design requirements into the Civil design to support entitlement, platting, and development approvals.",
        "These permit applications will be submitted using the electronic permitting submittal system (web-based system) for the respective jurisdictions where applicable."
    ],
    '210': [
        "Kimley-Horn will be available to provide miscellaneous project support at the direction of the Client. This task may include design meetings, additional permit support, permit research, or other miscellaneous tasks associated with the initial and future development of the project site. This task will also cover tasks such as design coordination meetings, scheduling, coordination with other client consultants, responses to additional rounds of agency comments."
    ]
}

# ============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ============================================================================

def add_title_section(doc, project_name, ipo_number):
    """Add title section at top of document."""
    
    # Main title - BOLD, CENTERED, 12pt
    para = doc.add_paragraph()
    run = para.add_run(project_name.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # IPO subtitle - BOLD, CENTERED, 10pt
    para = doc.add_paragraph()
    run = para.add_run(f'INDIVIDUAL PROJECT ORDER NUMBER {ipo_number}')
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Blank lines
    doc.add_paragraph()
    doc.add_paragraph()


def add_opening_paragraph(doc, client_name, master_agreement_date):
    """Add opening paragraph with master agreement reference."""
    
    para = doc.add_paragraph()
    run = para.add_run(
        f'Describing a specific agreement between Kimley-Horn and Associates, Inc. '
        f'(the Consultant), and {client_name} (the Client) in accordance with the terms of the '
        f'Master Agreement for Continuing Professional Services dated {master_agreement_date}, '
        f'which is incorporated herein by reference.'
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Blank line
    doc.add_paragraph()


def add_project_identification(doc, project_name, project_name_line2, project_manager, project_number):
    """Add Identification of Project section."""
    
    # Section heading - BOLD + UNDERLINED
    para = doc.add_paragraph()
    run = para.add_run('Identification of Project:')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.underline = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Blank line
    doc.add_paragraph()
    
    # Project Name - BOLD, with tab alignment at 2.5"
    para = doc.add_paragraph()
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
    
    run = para.add_run('Project Name:\t')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = para.add_run(project_name)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Second line of project name (if provided)
    if project_name_line2:
        para = doc.add_paragraph()
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
        
        run = para.add_run('\t')
        run = para.add_run(project_name_line2)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    # KH Project Manager
    para = doc.add_paragraph()
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
    
    run = para.add_run('KH Project Manager:\t')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = para.add_run(project_manager)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Project Number
    para = doc.add_paragraph()
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.LEFT)
    
    run = para.add_run('Project Number:\t')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = para.add_run(project_number)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Blank line
    doc.add_paragraph()


def add_section_with_heading(doc, heading_text, content_paragraphs):
    """Add a section with BOLD+UNDERLINED heading and content paragraphs."""
    
    # Section heading
    para = doc.add_paragraph()
    run = para.add_run(heading_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.underline = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Blank line
    doc.add_paragraph()
    
    # Content paragraphs
    for content in content_paragraphs:
        para = doc.add_paragraph()
        run = para.add_run(content)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        # Blank line after paragraph
        doc.add_paragraph()


def add_scope_of_services(doc, selected_tasks):
    """Add Specific scope of basic Services section with tasks."""
    
    # Section heading
    para = doc.add_paragraph()
    run = para.add_run('Specific scope of basic Services:')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.underline = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Blank line
    doc.add_paragraph()
    
    # Sub-section keywords for italic formatting
    sub_section_keywords = ['cover sheet', 'utility plan', 'site layout', 'site plan',
                           'grading plan', 'drainage plan', 'paving', 'erosion control',
                           'detail', 'existing conditions', 'demolition']
    
    # Add each selected task
    for task_num in sorted(selected_tasks.keys()):
        task = selected_tasks[task_num]
        descriptions = TASK_DESCRIPTIONS[task_num]
        
        # Task heading - BOLD + UNDERLINED
        para = doc.add_paragraph()
        run = para.add_run(f'Task {task_num} – {task["name"].replace("Civil ", "")}')
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.underline = True
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        # Blank line after task heading
        doc.add_paragraph()
        
        # Task description paragraphs
        for desc in descriptions:
            para = doc.add_paragraph()
            
            # Check if sub-section heading
            is_subsection = (len(desc) < 100 and 
                           any(kw in desc.lower() for kw in sub_section_keywords) and
                           not desc.endswith('.'))
            
            run = para.add_run(desc)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            
            # Apply ITALIC for sub-section headings
            if is_subsection:
                run.font.italic = True
            
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            
            # Only add blank line if NOT a sub-section heading
            if not is_subsection:
                doc.add_paragraph()


def add_simple_footer(doc, footer_text="rev 07/2024"):
    """Add simple footer text."""
    section = doc.sections[0]
    footer = section.footer
    
    para = footer.paragraphs[0]
    para.text = footer_text
    run = para.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def generate_msa_ipo_document(project_info, client_info, selected_tasks, output_path):
    """Generate Master Services Agreement IPO document."""
    
    doc = Document()
    
    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Add document sections
    add_title_section(doc, project_info['title'], project_info['ipo_number'])
    add_opening_paragraph(doc, client_info['name'], client_info['master_agreement_date'])
    add_project_identification(
        doc,
        project_info['name'],
        project_info.get('name_line2'),
        project_info['project_manager'],
        project_info['project_number']
    )
    
    # Overall Project Understanding
    if project_info.get('overall_understanding'):
        add_section_with_heading(
            doc,
            'Overall Project Understanding:',
            [project_info['overall_understanding']]
        )
    
    # Lot Specific Project Understanding
    if project_info.get('lot_understanding'):
        add_section_with_heading(
            doc,
            'Lot Specific Project Understanding:',
            [project_info['lot_understanding']]
        )
    
    # Scope of Services with tasks
    add_scope_of_services(doc, selected_tasks)
    
    # Add simple footer
    add_simple_footer(doc)
    
    # Save document
    doc.save(output_path)
    return output_path

# ============================================================================
# STREAMLIT APP
# ============================================================================

# Page configuration
st.set_page_config(
    page_title="MSA IPO Generator",
    page_icon="📋",
    layout="wide"
)

# Title
st.title("📋 Master Services Agreement - Individual Project Order Generator")
st.markdown("Generate IPO documents for projects under the Master Services Agreement")
st.markdown("---")

# Section 1: Document Header Information
st.header("📄 Document Header")
col1, col2 = st.columns(2)

with col1:
    project_title = st.text_input(
        "Project Title (appears at top) *",
        placeholder="e.g., USF Fletcher District – Phase 1",
        help="This appears as the main title at the top of the document"
    )
    
    project_name = st.text_input(
        "Project Name (Identification section) *",
        placeholder="e.g., USF Fletcher District – Phase 1",
        help="This appears in the 'Identification of Project' section"
    )

with col2:
    ipo_number = st.text_input(
        "IPO Number *",
        placeholder="e.g., 01",
        help="Individual Project Order number"
    )
    
    project_name_line2 = st.text_input(
        "Project Name Line 2",
        placeholder="e.g., Lot A Hotel, Conference Center and Retail",
        help="Optional second line for project name (will be aligned)"
    )

st.markdown("---")

# Section 2: Client Information
st.header("🏢 Client Information")
col3, col4 = st.columns(2)

with col3:
    client_name = st.text_input(
        "Client Name *",
        placeholder="e.g., ACE Fletcher LLC"
    )

with col4:
    master_agreement_date = st.text_input(
        "Master Agreement Date *",
        placeholder="e.g., August 15, 2024",
        help="Date of the Master Services Agreement"
    )

st.markdown("---")

# Section 3: Project Details
st.header("📍 Project Details")
col5, col6 = st.columns(2)

with col5:
    project_manager = st.text_input(
        "KH Project Manager *",
        placeholder="e.g., Dustin Ballard, PE"
    )

with col6:
    project_number = st.text_input(
        "Project Number *",
        placeholder="e.g., 145683001"
    )

overall_understanding = st.text_area(
    "Overall Project Understanding *",
    placeholder="Enter the overall project description...",
    height=100,
    help="High-level description of the overall project"
)

lot_understanding = st.text_area(
    "Lot Specific Project Understanding *",
    placeholder="Enter the lot-specific project description...",
    height=100,
    help="Description specific to this lot/phase"
)

st.markdown("---")

# Section 4: Task Selection and Fees
st.header("✅ Specific Scope of Basic Services")
st.markdown("Select the tasks to include in this IPO and enter the fee for each task.")

# Dictionary to store selected tasks
selected_tasks = {}

# Display each task with checkbox and fee input
for task_num in sorted(DEFAULT_FEES.keys()):
    task = DEFAULT_FEES[task_num]
    
    with st.expander(f"Task {task_num}: {task['name']}", expanded=True):
        col_check, col_fee = st.columns([3, 1])
        
        with col_check:
            task_selected = st.checkbox(
                f"Include Task {task_num}",
                key=f"check_{task_num}"
            )
        
        with col_fee:
            fee_amount = st.number_input(
                "Fee Amount ($)",
                min_value=0,
                value=None,
                placeholder=f"${task['amount']:,}",
                key=f"fee_{task_num}",
                disabled=not task_selected
            )
        
        st.markdown(f"**Fee Type:** {task['type']}")
        
        if task_selected:
            final_fee = fee_amount if fee_amount is not None else task['amount']
            selected_tasks[task_num] = {
                'name': task['name'],
                'fee': final_fee,
                'type': task['type']
            }

st.markdown("---")

# Section 5: Selected Tasks Summary
if selected_tasks:
    st.header("📊 Selected Tasks Summary")
    
    total_fee = 0
    for task_num in sorted(selected_tasks.keys()):
        task = selected_tasks[task_num]
        st.write(f"✓ Task {task_num}: {task['name']} — **${task['fee']:,}**")
        total_fee += task['fee']
    
    st.markdown("---")
    st.markdown(f"### **Total Fee: ${total_fee:,}**")
    st.markdown("---")
else:
    st.info("👆 Select at least one task to generate an IPO document")

# Section 6: Generate Button
st.header("📄 Generate IPO Document")

# Validation
required_fields = {
    'Project Title': project_title,
    'IPO Number': ipo_number,
    'Project Name': project_name,
    'Client Name': client_name,
    'Master Agreement Date': master_agreement_date,
    'KH Project Manager': project_manager,
    'Project Number': project_number,
    'Overall Project Understanding': overall_understanding,
    'Lot Specific Project Understanding': lot_understanding
}

missing_fields = [field for field, value in required_fields.items() if not value]

if missing_fields:
    st.warning(f"⚠️ Please fill in the following required fields: {', '.join(missing_fields)}")

if not selected_tasks:
    st.warning("⚠️ Please select at least one task")

# Generate button
can_generate = not missing_fields and bool(selected_tasks)

if st.button(
    "🚀 Generate IPO Document",
    type="primary",
    disabled=not can_generate
):
    with st.spinner("Generating IPO document..."):
        try:
            # Prepare data
            project_info = {
                'title': project_title,
                'ipo_number': ipo_number,
                'name': project_name,
                'name_line2': project_name_line2 if project_name_line2 else None,
                'project_manager': project_manager,
                'project_number': project_number,
                'overall_understanding': overall_understanding,
                'lot_understanding': lot_understanding
            }
            
            client_info = {
                'name': client_name,
                'master_agreement_date': master_agreement_date
            }
            
            # Generate document
            buffer = BytesIO()
            temp_path = '/tmp/temp_ipo.docx'
            generate_msa_ipo_document(project_info, client_info, selected_tasks, temp_path)
            
            # Read file into buffer
            with open(temp_path, 'rb') as f:
                buffer.write(f.read())
            buffer.seek(0)
            
            # Create filename
            filename = f"IPO_{ipo_number}_{project_name.replace(' ', '_')[:30]}.docx"
            
            st.success("✅ **Document Generated Successfully!**")
            
            # Download button
            st.download_button(
                label="📥 Download Word Document",
                data=buffer.getvalue(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
        except Exception as e:
            st.error(f"❌ **Error generating document:** {str(e)}")
            with st.expander("Show Error Details"):
                st.exception(e)

# Footer
st.markdown("---")
st.caption("MSA IPO Generator | Kimley-Horn Development Services")
